from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path("/Users/blackmamba/Documents/GitHub/winampGera")
OUTPUT = ROOT / "output" / "doc" / "winampgera-6-week-roadmap.docx"


def set_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def build_doc() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    set_base_styles(document)
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    document.add_heading("WinampGera 6-Week Roadmap", level=0)
    document.add_paragraph(
        "Scope basis: repository state as of March 13, 2026. "
        "This roadmap reflects the current Tkinter desktop app, the single-file "
        "application structure in winampgera.py, local JSON metrics, mocked unit "
        "tests, and the repo's current dependency and packaging gaps."
    )

    document.add_heading("Current Baseline", level=1)
    baseline_items = [
        "The app already supports opening and playing MP3, WAV, and FLAC files with python-vlc.",
        "The user interface is a fixed-size Tkinter desktop window with playback, stop, volume, and share controls.",
        "Share analytics are stored locally in winampgera_metrics.json rather than a service-backed telemetry pipeline.",
        "Most business logic, UI code, analytics, and playback orchestration live in a single file: winampgera.py.",
        "Tests exist and currently pass, but they focus on imports, mocked behavior, and basic regression checks rather than end-to-end playback or packaging flows.",
        "The repo does not indicate release packaging, installer automation, CI configuration, playlist management, or a persistent music library yet.",
    ]
    for item in baseline_items:
        add_bullet(document, item)

    document.add_heading("Milestones", level=1)
    milestone_items = [
        "Milestone 1: Core playback is reliable enough for repeated local use and easier to debug.",
        "Milestone 2: The codebase is split into testable modules so new features do not keep expanding one file.",
        "Milestone 3: WinampGera moves from single-track playback toward a usable personal music-player workflow.",
        "Milestone 4: The app is ready for a small beta release with clearer installation and support expectations.",
    ]
    for item in milestone_items:
        add_bullet(document, item)

    document.add_heading("Recommended Improvements to Fold Into This Roadmap", level=1)
    recommended_items = [
        "Modularize the app by extracting UI, playback, analytics, and app-state responsibilities out of winampgera.py.",
        "Expand test coverage beyond mocked unit tests to cover open, play, stop, share, and shutdown flows more realistically.",
        "Improve VLC and runtime dependency handling with OS-specific installation guidance and clearer failure states.",
        "Add a queue or playlist, recent-files support, and track metadata so the app becomes useful for repeated listening sessions.",
        "Add keyboard shortcuts plus playback progress and seek controls to make the desktop player feel faster and more complete.",
        "Refine the share experience with stronger now-playing copy, better feedback states, and safer analytics-file handling.",
        "Prepare CI and packaging early so a beta can be distributed without relying on manual environment setup.",
    ]
    for item in recommended_items:
        add_bullet(document, item)

    document.add_heading("Week-by-Week Plan", level=1)
    rows = [
        [
            "Week 1",
            "Stabilize the existing player",
            "Harden file loading, playback state transitions, dependency detection, and user-facing error handling around Tkinter and VLC.",
            "A cleaned playback-state checklist, better runtime error messages, a manual smoke-test script, and updated README troubleshooting notes.",
            "Requires working VLC installs on target OSes; risk that system-level VLC differences expose bugs that mocks do not catch.",
        ],
        [
            "Week 2",
            "Reduce architecture risk",
            "Split winampgera.py into focused modules such as UI, player service, analytics, and application state without changing the product surface yet.",
            "A modular package layout, preserved current behavior, and expanded tests around state transitions and analytics persistence.",
            "Depends on Week 1 stability baseline; risk of regressions during refactor because current tests do not cover full GUI flows.",
        ],
        [
            "Week 3",
            "Add music-library foundations",
            "Introduce a lightweight queue or playlist model, recent-files support, and basic track metadata display to make the app useful beyond one-off file opens.",
            "A first-pass playlist or queue experience, persisted recent tracks, metadata display for title, artist, or duration when available, and a UI update plan that fits the current fixed-window design.",
            "Depends on modular app state from Week 2; risk that Tkinter layout constraints make new controls feel cramped inside the current 600x400 window.",
        ],
        [
            "Week 4",
            "Improve engagement features",
            "Expand share and listening feedback with richer now-playing copy, clearer success states, keyboard shortcuts, and safer local metrics handling.",
            "A refined share flow, shortcut support for open, play or pause, and stop, metrics schema versioning or migration handling, and tests covering analytics-file corruption and recovery.",
            "Depends on Week 2 analytics extraction; risk that metrics remain too local to support broader product decisions without a future sync plan.",
        ],
        [
            "Week 5",
            "Prepare distribution and contributor workflow",
            "Add repeatable setup and release steps, including packaging investigation, pinned environment guidance, automated test execution in CI, and playback progress or seek feasibility work.",
            "A release checklist, CI test workflow, packaging spike results for at least one OS, clearer install instructions for VLC and Python, and a scoped plan for progress-bar or seek support in the current UI.",
            "Depends on a stable module layout and test suite; risk that desktop packaging is slower than expected because VLC bundling differs by platform.",
        ],
        [
            "Week 6",
            "Ship a small beta and plan the next cycle",
            "Run a controlled beta with a short feedback loop, fix the highest-severity issues, and turn findings into the next roadmap backlog.",
            "A beta build or runnable package, a feedback summary, a prioritized post-beta backlog, and a go or no-go recommendation for wider release.",
            "Depends on Week 5 release prep; risk that real-user issues cluster around platform setup rather than in-app UX, which can delay broader rollout.",
        ],
    ]

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Week", "Milestone", "Weekly Goal", "Deliverables", "Dependencies and Risks"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    document.add_heading("Cross-Cutting Dependencies", level=1)
    dependency_items = [
        "VLC remains the most important external dependency; release work should assume OS-specific install and packaging differences.",
        "Tkinter is suitable for the current desktop scope, but it constrains layout flexibility for bigger playlist, library, or visualization ideas.",
        "The single-file starting point makes refactoring a prerequisite for faster feature work; otherwise every new feature increases regression risk.",
        "Current tests are helpful for safe refactors, but they need GUI-adjacent and integration-style coverage before packaging or beta distribution.",
    ]
    for item in dependency_items:
        add_bullet(document, item)

    document.add_heading("Suggested Feature Backlog After Core Stabilization", level=1)
    backlog_items = [
        "Recent files so returning users can reopen tracks quickly.",
        "Track metadata display for title, artist, and duration.",
        "Keyboard shortcuts for open, play or pause, and stop.",
        "Playback progress and seek controls.",
        "Improved share messaging and better success-state UX.",
    ]
    for item in backlog_items:
        add_bullet(document, item)

    document.add_heading("Primary Risks", level=1)
    risk_items = [
        "Architecture risk: continuing to add features directly in winampgera.py will slow delivery and make debugging harder each week.",
        "Platform risk: desktop audio behavior and VLC availability can vary across macOS, Windows, and Linux.",
        "UX risk: the current window size and fixed layout may not absorb queue, playlist, and metadata features cleanly.",
        "Quality risk: passing unit tests may still miss playback, clipboard, file-dialog, and packaging issues seen only in real environments.",
        "Scope risk: adding major feature work before packaging and stability are improved could postpone a usable beta.",
    ]
    for item in risk_items:
        add_bullet(document, item)

    document.add_heading("Success Criteria at Week 6", level=1)
    success_items = [
        "A new contributor can install dependencies and run tests without guesswork.",
        "The app handles common playback and share flows without state confusion or silent failures.",
        "The codebase is modular enough that the next roadmap can add features without another large refactor first.",
        "There is a realistic beta artifact and a short, evidence-based backlog for the next release cycle.",
    ]
    for item in success_items:
        add_bullet(document, item)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
